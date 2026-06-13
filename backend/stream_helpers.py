"""Helper utilities for streaming handlers — audio transcription, attachments, file linking."""
import re
import time
from datetime import datetime, timedelta
from pathlib import Path
from urllib.parse import quote


_WEEKDAYS = ["Montag", "Dienstag", "Mittwoch", "Donnerstag", "Freitag", "Samstag", "Sonntag"]


def _owner_first() -> str:
    try:
        from identity import get_owner
    except ImportError:
        from backend.identity import get_owner
    return get_owner()["first_name"]

def build_time_context(last_user_ts: float | None = None) -> str:
    """Kompakter Zeitblock für jeden Agent-Aufruf.

    Liefert Datum, Uhrzeit, relative Tagesreferenzen und (falls bekannt) die
    Lücke zur letzten User-Nachricht. Ohne diesen Block muss das Modell aus
    Trainings-Cutoff oder Memory raten — beides geht regelmäßig schief.
    """
    now = datetime.now()
    today = now.date()
    wd = _WEEKDAYS[now.weekday()]
    iso_year, iso_week, _ = today.isocalendar()

    rel_lines = [
        f"- heute = {today.isoformat()} ({wd})",
        f"- gestern = {(today - timedelta(days=1)).isoformat()}",
        f"- vorgestern = {(today - timedelta(days=2)).isoformat()}",
        f"- morgen = {(today + timedelta(days=1)).isoformat()}",
    ]

    gap_line = ""
    if last_user_ts:
        try:
            last_dt = datetime.fromtimestamp(float(last_user_ts))
            delta = now - last_dt
            days = (today - last_dt.date()).days
            if days <= 0:
                if delta.total_seconds() < 3600:
                    gap = f"vor {int(delta.total_seconds() // 60)} min"
                else:
                    gap = f"vor {int(delta.total_seconds() // 3600)} h"
                gap_line = f"- letzte Nachricht in diesem Chat: {gap} (selber Tag)"
            elif days == 1:
                gap_line = f"- letzte Nachricht in diesem Chat: gestern ({last_dt.strftime('%H:%M')}), neuer Tag seitdem"
            elif days == 2:
                gap_line = f"- letzte Nachricht in diesem Chat: vorgestern ({last_dt.strftime('%H:%M')})"
            else:
                gap_line = f"- letzte Nachricht in diesem Chat: vor {days} Tagen ({last_dt.date().isoformat()})"
        except (ValueError, OSError):
            pass

    block = (
        "Zeitkontext (vom System, immer frisch):\n"
        f"- jetzt = {now.strftime('%Y-%m-%d %H:%M')} {wd}, KW {iso_week}, Zeitzone Europe/Berlin\n"
        + "\n".join(rel_lines)
    )
    if gap_line:
        block += "\n" + gap_line
    return block + "\n\n"


UPLOADS_DIR = Path(__file__).parent.parent / "data" / "uploads"
TEXT_EXTENSIONS = {'.txt', '.md', '.json', '.csv', '.py', '.js', '.ts', '.html', '.css', '.yml', '.yaml', '.toml', '.xml', '.sh', '.sql', '.env', '.log'}
AUDIO_EXTENSIONS = {'.m4a', '.mp3', '.ogg', '.wav', '.webm', '.flac', '.aac'}
PDF_MAX_CHARS = 60000
XLSX_MAX_CHARS_PER_SHEET = 20000
DOCX_MAX_CHARS = 60000

# Regex: matches absolute file paths in agent responses
_FILE_PATH_RE = re.compile(
    r'(?:^|\s)(/(?:Users|tmp|var|home|opt)[^\s\n\r\'"<>|*?]+\.(?:pdf|png|jpg|jpeg|gif|svg|webp|mp3|mp4|wav|m4a|csv|xlsx|xls|docx|doc|pptx|zip|tar\.gz|html|txt|md|json|py|js|ts))',
    re.MULTILINE
)


async def transcribe_audio(filepath: Path, name: str) -> str:
    """Transcribe audio file via Groq Whisper."""
    try:
        import httpx
        from server import _get_groq_key
        api_key = _get_groq_key()
        if not api_key:
            return f"[Sprachnachricht: {name} — kein Groq-Key konfiguriert]"
        data = filepath.read_bytes()
        ext = filepath.suffix.lower()
        mime_map = {'.m4a': 'audio/mp4', '.mp3': 'audio/mpeg', '.ogg': 'audio/ogg', '.wav': 'audio/wav', '.webm': 'audio/webm', '.flac': 'audio/flac', '.aac': 'audio/aac'}
        content_type = mime_map.get(ext, 'audio/mpeg')
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(
                "https://api.groq.com/openai/v1/audio/transcriptions",
                headers={"Authorization": f"Bearer {api_key}"},
                files={"file": (f"audio{ext}", data, content_type)},
                data={"model": "whisper-large-v3", "language": "de"},
            )
            result = resp.json()
            if resp.status_code != 200:
                return f"[Sprachnachricht: {name} — Transkription fehlgeschlagen]"
            transcript = result.get("text", "")
            return f"[Sprachnachricht: {name}]\n\n{transcript}"
    except Exception as e:
        return f"[Sprachnachricht: {name} — Fehler: {e}]"


def extract_pdf_text(filepath: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(filepath))
    chunks = []
    total = 0
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        chunks.append(f"--- Seite {i} ---\n{text}")
        total += len(text)
        if total >= PDF_MAX_CHARS:
            chunks.append(f"[… weitere {len(reader.pages) - i} Seiten wegen Länge abgeschnitten]")
            break
    return "\n\n".join(chunks) if chunks else "[PDF enthält keinen extrahierbaren Text]"


def extract_xlsx_text(filepath: Path) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(str(filepath), data_only=True, read_only=True)
    blocks = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_md = []
        header_written = False
        char_count = 0
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v).replace("|", "\\|").replace("\n", " ") for v in row]
            if not any(c.strip() for c in cells):
                continue
            line = "| " + " | ".join(cells) + " |"
            rows_md.append(line)
            if not header_written:
                rows_md.append("| " + " | ".join(["---"] * len(cells)) + " |")
                header_written = True
            char_count += len(line)
            if char_count >= XLSX_MAX_CHARS_PER_SHEET:
                rows_md.append("[… Sheet wegen Länge abgeschnitten]")
                break
        if rows_md:
            blocks.append(f"### Sheet: {sheet_name}\n" + "\n".join(rows_md))
    wb.close()
    return "\n\n".join(blocks) if blocks else "[XLSX enthält keine lesbaren Daten]"


def extract_docx_text(filepath: Path) -> str:
    from docx import Document
    doc = Document(str(filepath))
    parts = []
    total = 0
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style = (para.style.name or "") if para.style else ""
        if style.startswith("Heading"):
            level = style.replace("Heading", "").strip() or "1"
            try:
                prefix = "#" * max(1, min(6, int(level)))
            except ValueError:
                prefix = "##"
            parts.append(f"{prefix} {text}")
        else:
            parts.append(text)
        total += len(text)
        if total >= DOCX_MAX_CHARS:
            parts.append("[… weitere Absätze wegen Länge abgeschnitten]")
            break
    for i, table in enumerate(doc.tables, start=1):
        if total >= DOCX_MAX_CHARS:
            break
        rows_md = []
        for row_idx, row in enumerate(table.rows):
            cells = [(cell.text or "").replace("|", "\\|").replace("\n", " ").strip() for cell in row.cells]
            if not any(cells):
                continue
            rows_md.append("| " + " | ".join(cells) + " |")
            if row_idx == 0:
                rows_md.append("| " + " | ".join(["---"] * len(cells)) + " |")
        if rows_md:
            parts.append(f"\n**Tabelle {i}**\n" + "\n".join(rows_md))
            total += sum(len(r) for r in rows_md)
    return "\n\n".join(parts) if parts else "[DOCX enthält keinen extrahierbaren Text]"


async def build_attachment_context(attachments: list[dict]) -> str:
    """Read uploaded files and build context string to append to message."""
    if not attachments:
        return ""
    parts = []
    for att in attachments:
        url = att.get("url", "")
        name = att.get("name", "unknown")
        mime = att.get("type", "")
        if url.startswith("/uploads/"):
            filepath = UPLOADS_DIR / url.split("/uploads/", 1)[1]
        elif url.startswith("/api/serve"):
            from urllib.parse import unquote, urlparse, parse_qs
            qs = parse_qs(urlparse(url).query)
            filepath = Path(unquote(qs.get("path", [""])[0]))
        else:
            parts.append(f"[Datei: {name} — URL: {url}]")
            continue
        if not filepath.exists():
            parts.append(f"[Datei: {name} — nicht gefunden]")
            continue
        ext = filepath.suffix.lower()
        if ext in AUDIO_EXTENSIONS:
            transcript = await transcribe_audio(filepath, name)
            parts.append(transcript)
        elif ext in TEXT_EXTENSIONS:
            try:
                content = filepath.read_text(errors='replace')[:8000]
                parts.append(f"[Dateiinhalt: {name}]\n```\n{content}\n```")
            except Exception:
                parts.append(f"[Datei: {name} — Lesefehler]")
        elif ext == '.pdf':
            try:
                content = extract_pdf_text(filepath)
                parts.append(f"[PDF-Inhalt: {name}]\n\n{content}")
            except Exception as e:
                parts.append(f"[Datei: {name} — PDF-Lesefehler: {e}]")
        elif ext in {'.xlsx', '.xlsm'}:
            try:
                content = extract_xlsx_text(filepath)
                parts.append(f"[Tabelle: {name}]\n\n{content}")
            except Exception as e:
                parts.append(f"[Datei: {name} — XLSX-Lesefehler: {e}]")
        elif ext == '.docx':
            try:
                content = extract_docx_text(filepath)
                parts.append(f"[DOCX-Inhalt: {name}]\n\n{content}")
            except Exception as e:
                parts.append(f"[Datei: {name} — DOCX-Lesefehler: {e}]")
        elif ext == '.doc':
            parts.append(f"[Datei: {name} — altes .doc-Format wird nicht unterstützt, bitte als .docx speichern]")
        else:
            parts.append(f"[Datei: {name} — Typ: {mime}, Pfad: {filepath}]")
    if not parts:
        return ""
    return "\n\n" + "\n\n".join(parts)


def build_focus_snapshot() -> str:
    """Dünner Stand aus /fokus, Kalender, Inbox und Pipeline für jeden Chat-Send."""
    try:
        from modules.fokus.core import render_focus_system_snapshot
    except Exception:
        return ""
    return render_focus_system_snapshot()


def build_context_router_context(message: str, conv_id: str) -> str:
    """Automatischer Context Router: relevante vorhandene Quellen laden und loggen."""
    try:
        from context_router import build_context
        return build_context(message, conv_id, log_run=False)
    except Exception as e:
        try:
            import workflows
            run_id = workflows.start_run(
                "context.router",
                "Context Router",
                trigger="chat_turn",
                subject_type="conversation",
                subject_ref=conv_id,
                conversation_id=conv_id,
                input_data={"query": message[:1000]},
            )
            workflows.add_step(run_id, "router_error", "Router Fehler", "error", str(e), {})
            workflows.finish_run(run_id, "error", error=str(e))
        except Exception:
            pass
        return ""


def build_dreaming_context() -> str:
    """Leiser Dreaming-Kontext für Live-Antworten.

    Der Block ist bewusst klein und handlungsarm: Er soll Klaus' Ton und
    Priorisierung kalibrieren, aber der Nutzer keine weitere Lesepflicht geben.
    """
    try:
        import dreaming_module
        return dreaming_module.live_context()
    except Exception:
        return ""


def build_focus_item_context(conv_id: str) -> str:
    """Frischer Item-Kontext für Conversations, die an ein Fokus-Item gekoppelt sind.
    Liefert Fokus-Metadaten, Body und letzte Notizen, damit Klaus pro Turn weiß,
    woran konkret gearbeitet wird. Leerer String, wenn die Conv kein Fokus-Item ist."""
    if not conv_id:
        return ""
    try:
        from db import get_db
        from modules.fokus.core import _focus_resolve_identity, _build_fokus_line
    except Exception:
        return ""
    try:
        with get_db() as db:
            row = db.execute(
                "SELECT item_key FROM focus_item_conv WHERE conv_id = ?",
                (conv_id,),
            ).fetchone()
    except Exception:
        return ""
    if not row:
        return ""
    item_key = row[0]
    item = _focus_resolve_identity(item_key=item_key)
    if not item or item.get("status") != "open":
        return ""
    title = item.get("title") or ""
    line = _build_fokus_line(
        item.get("bucket") or "later",
        title,
        item.get("date") or "",
        item.get("people") or [],
        item.get("projects") or [],
        item.get("tags") or [],
    )
    body_text = (item.get("body") or "").strip()
    # Notizen aus focus_updates
    notes: list[str] = []
    try:
        with get_db() as db:
            rows = db.execute(
                "SELECT text, source, ts FROM focus_updates WHERE item_key = ? ORDER BY ts ASC",
                (item_key,),
            ).fetchall()
            for r in rows:
                src = r[1] or "text"
                txt = (r[0] or "").strip().replace("\n", " ")
                if len(txt) > 400:
                    txt = txt[:400] + "…"
                notes.append(f"- ({src}) {txt}")
    except Exception:
        pass

    parts = [
        "Fokus-Item-Kontext (vom System, immer frisch — Quelle der Wahrheit, NICHT aus Chat-History rekonstruieren):",
        f"Titel: {title}",
    ]
    if line:
        parts.append(f"Fokus-Zeile: {line}")
    if body_text:
        if len(body_text) > 1200:
            body_text = body_text[:1200] + "…"
        parts.append("Body / Notizen aus /fokus:\n" + body_text)
    if notes:
        parts.append("Sprach-/Text-Notizen zu diesem Item:\n" + "\n".join(notes[-10:]))
    parts.append(
        "Auftrag: Du bist hier in einem fokussierten Mini-Thread genau zu diesem Item. "
        "Antworte konkret darauf bezogen. Wenn der Stand sich ändert (Termin gemacht, "
        "erledigt, neue Person, neues Datum), aktualisiere das Item über die Fokus-API "
        "und sag in einem Satz, was du geändert hast."
    )
    parts.append(
        "\nSlot-Anlage — Pflicht bei konkreter Zeitangabe:\n"
        f"Wenn {_owner_first()} Datum UND Uhrzeit nennt (z.B. \"26.05. 14 bis 16 Uhr\", "
        "\"morgen 9:30 für eine Stunde\"), MUSST du zusätzlich einen Kalender-Slot "
        "anlegen — sonst erscheint das Item als Ganztageintrag statt als Block. "
        "Endpoint:\n"
        "  curl -s -X POST http://127.0.0.1:8890/api/fokus/slots \\\n"
        "       -H 'Content-Type: application/json' \\\n"
        f"       -d '{{\"title\":\"{title}\",\"day_iso\":\"YYYY-MM-DD\","
        "\"start_min\":<Minuten ab Mitternacht>,\"dur_min\":<Dauer in Minuten>}}'\n"
        "Umrechnung: \"14 bis 16 Uhr\" → start_min=840, dur_min=120. "
        "\"9:30 eine Stunde\" → start_min=570, dur_min=60. "
        "title MUSS exakt der oben genannte Item-Titel sein, damit der Slot zum Item gemappt wird."
    )
    return "\n".join(parts) + "\n\n"


def build_session_context(conv_id: str) -> str:
    """Sagt Klaus, in welcher Chat-Session er gerade sitzt.

    Ohne diesen Block weiss Klaus nicht, welchen Pane er bedient — Aufträge wie
    "benenn diesen Chat um" landen sonst in einer anderen Session.
    """
    if not conv_id:
        return ""
    title = ""
    try:
        from db import get_db
        with get_db() as db:
            row = db.execute(
                "SELECT title FROM conversations WHERE id = ?",
                (conv_id,),
            ).fetchone()
        if row and row[0]:
            title = str(row[0])
    except Exception:
        pass
    lines = [
        "Aktuelle Chat-Session (vom System, immer frisch):",
        f"- conversationId = {conv_id}",
    ]
    if title:
        lines.append(f"- aktueller Titel = {title}")
    lines.append(
        "- \"dieser Chat\" / \"diese Session\" meint IMMER genau diese conversationId, "
        f"nie eine andere. Wenn {_owner_first()} sagt \"benenn diesen Chat um\", rufst du auf:"
    )
    lines.append(
        f"  curl -s -X PATCH http://127.0.0.1:8890/api/conversations/{conv_id} "
        "-H 'Content-Type: application/json' -d '{\"title\":\"<neuer Titel>\"}'"
    )
    return "\n".join(lines) + "\n\n"


def build_person_context(message: str, conv_id: str = "", max_people: int = 3) -> str:
    """Zieht Person-Steckbriefe in den Prompt, sobald jemand erwähnt wird.

    Nutzt entities.detect_people auf die aktuelle User-Nachricht und liest
    person_summary.body für die top-N Treffer. Ohne Treffer leerer String.
    """
    if not message:
        return ""
    try:
        from backend import entities
        import sqlite3
        from pathlib import Path
    except Exception:
        return ""
    try:
        people_ids = entities.detect_people(message)
    except Exception:
        return ""
    if not people_ids:
        return ""
    db_path = Path(__file__).parent.parent / "data" / "people.db"
    rows: list[tuple[str, str]] = []
    try:
        conn = sqlite3.connect(db_path, timeout=5)
        conn.row_factory = sqlite3.Row
        placeholders = ",".join("?" * len(people_ids))
        cur = conn.execute(
            f"SELECT p.name AS name, s.body AS body "
            f"FROM people p LEFT JOIN person_summary s ON s.person_id=p.id "
            f"WHERE p.id IN ({placeholders}) "
            f"ORDER BY s.updated_at DESC NULLS LAST",
            people_ids,
        )
        for r in cur.fetchall():
            body = (r["body"] or "").strip()
            if not body:
                continue
            rows.append((r["name"], body))
            if len(rows) >= max_people:
                break
        conn.close()
    except Exception:
        return ""
    if not rows:
        return ""
    parts = [
        "Kontext zu erwaehnten Personen (vom System, aus people.db/person_summary, "
        "verdichtet aus allen Quellen):",
    ]
    for name, body in rows:
        parts.append(f"\n## {name}\n{body}")
    return "\n".join(parts) + "\n\n"


def build_focus_quick_add_context(conv_id: str) -> str:
    """Quick-Add-Modus: pro Turn frischer Kontext für Anlegen neuer Fokus-Items."""
    if not conv_id:
        return ""
    try:
        from db import get_db
        from modules.fokus.core import _focus_list_items
    except Exception:
        return ""
    try:
        with get_db() as db:
            row = db.execute(
                "SELECT 1 FROM focus_quick_add_conv WHERE conv_id = ?",
                (conv_id,),
            ).fetchone()
    except Exception:
        return ""
    if not row:
        return ""
    excerpt_lines: list[str] = []
    try:
        items = _focus_list_items(include_done=False, business_only=False)[:20]
        for it in items:
            label = it.get("date") or "ohne Datum"
            excerpt_lines.append(f"- [{it.get('bucket')}] {label} — {it.get('title')}")
    except Exception:
        excerpt_lines = []
    parts = [
        "Fokus-Quick-Add-Kontext (vom System, immer frisch):",
        f"Auftrag: Du hilfst {_owner_first()}, ein NEUES Item in `/fokus` anzulegen. "
        "Frag knapp nach was fehlt (Wann? Wer? Worum geht's?). Wenn klar genug, "
        "lege es per POST auf `/api/fokus` im passenden Bucket an. "
        "Bestätige in einem Satz, was du angelegt hast. "
        "Halte die Konversation kurz und handlungsorientiert.",
        "",
        "Slot-Anlage — Pflicht bei konkreter Zeitangabe:",
        f"Wenn {_owner_first()} Datum UND Uhrzeit nennt (z.B. \"morgen 14 bis 16 Uhr\", "
        "\"Mittwoch 9:30\", \"heute Nachmittag um 15 Uhr\"), MUSST du zusätzlich "
        "einen Kalender-Slot anlegen über:",
        "  curl -s -X POST http://127.0.0.1:8890/api/fokus/slots \\",
        "       -H 'Content-Type: application/json' \\",
        "       -d '{\"title\":\"<exakt der Item-Titel den du gerade angelegt hast>\","
        "\"day_iso\":\"YYYY-MM-DD\",\"start_min\":<Minuten ab Mitternacht>,\"dur_min\":<Dauer in Minuten>}'",
        "Beispiele für die Umrechnung:",
        "  \"14 bis 16 Uhr\" → start_min=840, dur_min=120",
        "  \"9:30 für eine Stunde\" → start_min=570, dur_min=60",
        "  \"15 Uhr\" (Dauer offen) → start_min=900, dur_min=30",
        "Niemals einen Ganztagsblock anlegen, wenn eine Uhrzeit genannt wurde. "
        "Wenn nur Datum, kein Slot.",
        "",
        "Aktueller Fokus-Auszug (zur Orientierung):",
        "\n".join(excerpt_lines) if excerpt_lines else "- derzeit keine offenen Items",
    ]
    return "\n".join(parts) + "\n\n"


def build_whatsapp_project_context(project_id: str, per_chat_limit: int = 12) -> str:
    """Liefert die letzten Nachrichten aller WhatsApp-Chats, die an dieses Projekt gepinnt sind.

    Wird als Block in den initialen Stream-Kontext geschoben, damit Klaus in
    Projekt-Chats automatisch weiss, was per WhatsApp lief, ohne dass der Nutzer
    es nachreichen muss.
    """
    if not project_id:
        return ""
    try:
        from server import _wa_chats_for_project, _wa_db, _wa_has_col
    except Exception:
        return ""
    chats = _wa_chats_for_project(project_id)
    if not chats:
        return ""
    blocks = []
    with _wa_db() as con:
        has_summary = _wa_has_col(con, "messages", "summary")
        for c in chats:
            cid = c["chat_id"]
            name = c["name"]
            extra = ", summary" if has_summary else ""
            rows = con.execute(
                f"SELECT from_me, ts, type, body, transcript{extra} FROM messages "
                f"WHERE chat_id=? ORDER BY ts DESC LIMIT ?",
                (cid, per_chat_limit),
            ).fetchall()
            if not rows:
                continue
            lines = []
            for r in reversed(rows):
                who = "Du" if r["from_me"] else name
                txt = r["transcript"] or r["body"] or ""
                if not txt:
                    if has_summary and r["summary"]:
                        txt = r["summary"]
                    else:
                        t = r["type"] or ""
                        txt = f"[{t}]" if t else ""
                txt = (txt or "").replace("\n", " ").strip()
                if len(txt) > 240:
                    txt = txt[:240] + "…"
                lines.append(f"  {who}: {txt}")
            header = f"WhatsApp-Chat „{name}“ (letzte {len(lines)}):"
            blocks.append(header + "\n" + "\n".join(lines))
    if not blocks:
        return ""
    return (
        "An dieses Projekt gepinnte WhatsApp-Chats (Quelle: WhatsApp, nur Verweis):\n\n"
        + "\n\n".join(blocks)
        + "\n\n"
    )


def linkify_file_paths(text: str) -> str:
    """Replace local file paths in agent output with download links."""
    def _replace(m: re.Match) -> str:
        path = m.group(1)
        if not Path(path).exists():
            return m.group(0)
        name = Path(path).name
        url = f"/api/serve?path={quote(path, safe='')}"
        prefix = m.group(0)[:-len(m.group(1))]
        return f'{prefix}[{name}]({url})'
    return _FILE_PATH_RE.sub(_replace, text)


def update_letzter_stand(project: str, agent: str, response: str, projects_roots):
    """Append a line to the 'Letzter Stand' section of the project's CLAUDE.md."""
    if not project:
        return
    p = Path(project)
    if not p.exists():
        p = projects_roots[0] / project
    claude_md = p / "CLAUDE.md"
    if not claude_md.exists():
        return
    try:
        content = claude_md.read_text()
    except (UnicodeDecodeError, OSError):
        return
    summary = response.strip().split('\n')[0][:80]
    if len(response.strip().split('\n')[0]) > 80:
        summary += '...'
    date = time.strftime('%Y-%m-%d')
    new_line = f"{date} {agent} — {summary}"

    if '## Letzter Stand' in content:
        parts = content.split('## Letzter Stand')
        after = parts[1]
        lines = [l for l in after.strip().split('\n') if l.strip()]
        lines = [new_line] + lines[:4]
        content = parts[0] + '## Letzter Stand\n' + '\n'.join(lines) + '\n'
    else:
        content = content.rstrip() + f'\n\n## Letzter Stand\n{new_line}\n'

    try:
        claude_md.write_text(content)
    except OSError:
        pass
